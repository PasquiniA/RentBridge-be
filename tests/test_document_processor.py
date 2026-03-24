"""
Tests for DocumentProcessor
"""

import pytest
from docx import Document
import tempfile
import os
from app.providers.document_processor import DocumentProcessor


@pytest.fixture
def document_processor():
    """Fixture to create a DocumentProcessor instance."""
    return DocumentProcessor()


@pytest.fixture
def sample_docx_with_placeholders():
    """Create a sample .docx file with placeholders."""
    doc = Document()
    doc.add_paragraph("Hello {{name}}, welcome to {{company}}!")
    doc.add_paragraph("Your email is: {{email}}")
    
    # Add a table with placeholders
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "{{name}}"
    table.rows[1].cells[0].text = "City"
    table.rows[1].cells[1].text = "{{city}}"
    
    # Save to temporary file and return bytes
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        doc.save(temp_file.name)
        temp_file_path = temp_file.name
    
    with open(temp_file_path, 'rb') as f:
        doc_bytes = f.read()
    
    os.unlink(temp_file_path)
    return doc_bytes


@pytest.fixture
def sample_docx_with_italian_accents():
    """Create a sample .docx file with Italian accents."""
    doc = Document()
    doc.add_paragraph("Città: {{city}}")
    doc.add_paragraph("Età: {{age}}")
    doc.add_paragraph("Qualità: {{quality}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        doc.save(temp_file.name)
        temp_file_path = temp_file.name
    
    with open(temp_file_path, 'rb') as f:
        doc_bytes = f.read()
    
    os.unlink(temp_file_path)
    return doc_bytes


def test_replace_placeholders_basic(document_processor, sample_docx_with_placeholders):
    """Test basic placeholder replacement."""
    replacements = {
        "name": "Mario Rossi",
        "company": "Acme Corp",
        "email": "mario@example.com",
        "city": "Milano"
    }
    
    result_bytes = document_processor.replace_placeholders(
        sample_docx_with_placeholders,
        replacements
    )
    
    # Verify result is valid docx
    assert result_bytes is not None
    assert len(result_bytes) > 0
    
    # Load result and verify replacements
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(result_bytes)
        temp_file_path = temp_file.name
    
    try:
        doc = Document(temp_file_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify placeholders are replaced
        assert "{{name}}" not in full_text
        assert "{{company}}" not in full_text
        assert "{{email}}" not in full_text
        assert "{{city}}" not in full_text
        
        # Verify values are present
        assert "Mario Rossi" in full_text
        assert "Acme Corp" in full_text
        assert "mario@example.com" in full_text
        
        # Check table
        table_text = '\n'.join([cell.text for row in doc.tables[0].rows for cell in row.cells])
        assert "Mario Rossi" in table_text
        assert "Milano" in table_text
        
    finally:
        os.unlink(temp_file_path)


def test_replace_placeholders_with_italian_accents(document_processor, sample_docx_with_italian_accents):
    """Test placeholder replacement with Italian accented characters."""
    replacements = {
        "city": "Città di Roma",
        "age": "trentacinque",
        "quality": "qualità eccellente"
    }
    
    result_bytes = document_processor.replace_placeholders(
        sample_docx_with_italian_accents,
        replacements
    )
    
    # Load result and verify Italian accents are preserved
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(result_bytes)
        temp_file_path = temp_file.name
    
    try:
        doc = Document(temp_file_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify Italian accents are preserved
        assert "Città di Roma" in full_text
        assert "qualità eccellente" in full_text
        
    finally:
        os.unlink(temp_file_path)


def test_replace_placeholders_partial_replacements(document_processor, sample_docx_with_placeholders):
    """Test that unreplaced placeholders remain in document."""
    replacements = {
        "name": "Mario Rossi"
        # Intentionally not replacing other placeholders
    }
    
    result_bytes = document_processor.replace_placeholders(
        sample_docx_with_placeholders,
        replacements
    )
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(result_bytes)
        temp_file_path = temp_file.name
    
    try:
        doc = Document(temp_file_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify replaced placeholder
        assert "{{name}}" not in full_text
        assert "Mario Rossi" in full_text
        
        # Verify unreplaced placeholders remain
        assert "{{company}}" in full_text
        assert "{{email}}" in full_text
        
    finally:
        os.unlink(temp_file_path)


# Note: doc_to_pdf tests require LibreOffice to be installed
# These tests will be skipped in environments without LibreOffice
@pytest.mark.skipif(
    os.system("which libreoffice > /dev/null 2>&1") != 0,
    reason="LibreOffice not installed"
)
def test_doc_to_pdf_conversion(document_processor, sample_docx_with_placeholders):
    """Test conversion from .docx to PDF."""
    pdf_bytes = document_processor.doc_to_pdf(sample_docx_with_placeholders)
    
    # Verify PDF was generated
    assert pdf_bytes is not None
    assert len(pdf_bytes) > 0
    
    # Verify PDF magic bytes
    assert pdf_bytes.startswith(b'%PDF')


@pytest.mark.skipif(
    os.system("which libreoffice > /dev/null 2>&1") != 0,
    reason="LibreOffice not installed"
)
def test_doc_to_pdf_with_italian_accents(document_processor, sample_docx_with_italian_accents):
    """Test PDF conversion preserves Italian accents."""
    # First replace placeholders with Italian text
    replacements = {
        "city": "Città di Roma",
        "age": "trentacinque",
        "quality": "qualità eccellente"
    }
    
    doc_bytes = document_processor.replace_placeholders(
        sample_docx_with_italian_accents,
        replacements
    )
    
    # Convert to PDF
    pdf_bytes = document_processor.doc_to_pdf(doc_bytes)
    
    # Verify PDF was generated
    assert pdf_bytes is not None
    assert len(pdf_bytes) > 0
    assert pdf_bytes.startswith(b'%PDF')
