"""
Document processing provider for mail-merge and PDF conversion.

This module handles:
- Placeholder replacement in .doc/.docx files using python-docx
- Conversion from .doc/.docx to PDF using LibreOffice
- UTF-8 support for Italian accents (à, è, ì, ò, ù)
"""

import os
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class DocumentProcessor:
    """
    Handles document processing operations including placeholder replacement
    and conversion to PDF format.
    """
    
    def __init__(self):
        """Initialize the DocumentProcessor."""
        pass
    
    def replace_placeholders(self, doc_bytes: bytes, replacements: Dict[str, str]) -> bytes:
        """
        Replace placeholders in a .docx document with provided values.
        
        Placeholders are in the format {{variable_name}}.
        
        Args:
            doc_bytes: Bytes of the .docx document
            replacements: Dictionary mapping placeholder names to replacement values
            
        Returns:
            Bytes of the modified .docx document with placeholders replaced
            
        Raises:
            ValueError: If document cannot be processed
        """
        # Create temporary file for input document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_input:
            temp_input.write(doc_bytes)
            temp_input_path = temp_input.name
        
        try:
            # Load the document
            doc = Document(temp_input_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        # Replace in runs to preserve formatting
                        self._replace_in_paragraph(paragraph, placeholder, value)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in paragraph.text:
                                    self._replace_in_paragraph(paragraph, placeholder, value)
            
            # Replace placeholders in headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, value)
                
                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, value)
            
            # Save modified document to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
                temp_output_path = temp_output.name
            
            doc.save(temp_output_path)
            
            # Read the modified document bytes
            with open(temp_output_path, 'rb') as f:
                modified_bytes = f.read()
            
            # Clean up temporary files
            os.unlink(temp_output_path)
            
            return modified_bytes
            
        finally:
            # Clean up input temporary file
            if os.path.exists(temp_input_path):
                os.unlink(temp_input_path)
    
    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str):
        """
        Replace placeholder in a paragraph while preserving formatting.
        
        This method handles the case where a placeholder might be split across
        multiple runs in the paragraph.
        
        Args:
            paragraph: The paragraph object to modify
            placeholder: The placeholder string to find (e.g., "{{name}}")
            value: The replacement value
        """
        # Get full text
        full_text = paragraph.text
        
        if placeholder not in full_text:
            return
        
        # Replace in full text
        new_text = full_text.replace(placeholder, value)
        
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""
        
        # Add new text to first run (or create new run if none exist)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
    
    def doc_to_pdf(self, doc_bytes: bytes) -> bytes:
        """
        Convert a .doc or .docx document to PDF using LibreOffice.
        
        This method uses LibreOffice in headless mode to perform the conversion.
        UTF-8 encoding is ensured for proper handling of Italian accents.
        
        Args:
            doc_bytes: Bytes of the .doc or .docx document
            
        Returns:
            Bytes of the generated PDF document
            
        Raises:
            RuntimeError: If conversion fails
            FileNotFoundError: If LibreOffice is not installed
        """
        # Create temporary directory for conversion
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)
            
            # Write input document to temporary file
            input_file = temp_dir_path / "input.docx"
            with open(input_file, 'wb') as f:
                f.write(doc_bytes)
            
            # Expected output PDF file
            output_file = temp_dir_path / "input.pdf"
            
            try:
                # Run LibreOffice conversion
                # --headless: Run without GUI
                # --convert-to pdf: Convert to PDF format
                # --outdir: Output directory
                result = subprocess.run(
                    [
                        'libreoffice',
                        '--headless',
                        '--convert-to',
                        'pdf',
                        '--outdir',
                        str(temp_dir_path),
                        str(input_file)
                    ],
                    capture_output=True,
                    text=True,
                    timeout=30,  # 30 second timeout
                    env={**os.environ, 'LC_ALL': 'en_US.UTF-8', 'LANG': 'en_US.UTF-8'}
                )
                
                # Check if conversion was successful
                if result.returncode != 0:
                    raise RuntimeError(
                        f"LibreOffice conversion failed with return code {result.returncode}. "
                        f"Error: {result.stderr}"
                    )
                
                # Check if output file was created
                if not output_file.exists():
                    raise RuntimeError(
                        f"PDF file was not created. LibreOffice output: {result.stdout}"
                    )
                
                # Read the generated PDF
                with open(output_file, 'rb') as f:
                    pdf_bytes = f.read()
                
                return pdf_bytes
                
            except FileNotFoundError:
                raise FileNotFoundError(
                    "LibreOffice is not installed or not found in PATH. "
                    "Please install LibreOffice to use document conversion."
                )
            except subprocess.TimeoutExpired:
                raise RuntimeError(
                    "LibreOffice conversion timed out after 30 seconds. "
                    "The document may be too large or complex."
                )
