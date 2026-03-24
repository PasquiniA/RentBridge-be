"""
Mail-Merge Service for document processing and PDF generation.

This service orchestrates the mail-merge workflow:
1. Replace placeholders in document template
2. Convert document to PDF
3. Upload PDF to Supabase Storage
4. Cleanup temporary files
"""

import os
import tempfile
import time
from typing import Dict, Any, Optional
from pathlib import Path
from pydantic import BaseModel, Field
from datetime import datetime

from app.providers.document_processor import DocumentProcessor
from app.providers.supabase_storage import SupabaseStorageClient, SupabaseUploadResult
from app.utils.exceptions import (
    InvalidTemplateError,
    MissingFieldsError,
    ConversionError,
    SupabaseUploadError
)


class MailMergeResult(BaseModel):
    """Result of a mail-merge operation."""
    
    pdf_bytes: bytes = Field(..., description="Bytes of the generated PDF")
    supabase_upload: SupabaseUploadResult = Field(
        ...,
        description="Information about the uploaded file in Supabase"
    )
    placeholders_replaced: int = Field(
        ...,
        ge=0,
        description="Number of placeholders replaced in the document"
    )
    processing_time_ms: float = Field(
        ...,
        ge=0,
        description="Time taken to process the request in milliseconds"
    )


class MailMergeService:
    """
    Service for orchestrating mail-merge operations.
    
    This service handles the complete workflow of:
    - Validating merge data against template placeholders
    - Replacing placeholders in document templates
    - Converting documents to PDF format
    - Uploading PDFs to Supabase Storage
    - Cleaning up temporary files
    """
    
    def __init__(
        self,
        storage_client: SupabaseStorageClient,
        document_processor: Optional[DocumentProcessor] = None
    ):
        """
        Initialize the MailMergeService.
        
        Args:
            storage_client: Client for uploading files to Supabase Storage
            document_processor: Processor for document operations (optional, creates default if None)
        """
        self.storage_client = storage_client
        self.document_processor = document_processor or DocumentProcessor()
    
    async def merge_and_upload_document(
        self,
        template_bytes: bytes,
        merge_data: Dict[str, Any],
        file_name: Optional[str] = None,
        user_id: Optional[str] = None
    ) -> MailMergeResult:
        """
        Execute complete mail-merge workflow.
        
        This method:
        1. Extracts placeholders from template
        2. Validates merge_data contains all required placeholders
        3. Replaces placeholders with values from merge_data
        4. Converts document to PDF
        5. Uploads PDF to Supabase Storage
        6. Cleans up temporary files
        
        Args:
            template_bytes: Bytes of the .doc/.docx template file
            merge_data: Dictionary with values to replace placeholders
            file_name: Optional custom filename for the PDF (without extension)
            user_id: Optional user ID for organizing files in Supabase
            
        Returns:
            MailMergeResult with PDF bytes, upload info, and processing metrics
            
        Raises:
            InvalidTemplateError: If template is not valid
            MissingFieldsError: If merge_data is missing required placeholders
            ConversionError: If document conversion fails
            SupabaseUploadError: If upload to Supabase fails
        """
        start_time = time.time()
        temp_files = []
        
        try:
            # Step 1: Extract placeholders from template
            placeholders = self.extract_placeholders(template_bytes)
            
            if not placeholders:
                raise InvalidTemplateError(
                    "Template does not contain any placeholders in {{variable}} format"
                )
            
            # Step 2: Validate merge_data contains all placeholders
            validation_result = self.validate_merge_data(placeholders, merge_data)
            
            if not validation_result["is_valid"]:
                missing = validation_result["missing_fields"]
                raise MissingFieldsError(
                    missing_fields=missing,
                    template_placeholders=placeholders
                )
            
            # Step 3: Replace placeholders in document
            # Convert all merge_data values to strings
            replacements = {key: str(value) for key, value in merge_data.items()}
            
            try:
                merged_doc_bytes = self.document_processor.replace_placeholders(
                    template_bytes,
                    replacements
                )
            except Exception as e:
                raise InvalidTemplateError(f"Failed to process template: {str(e)}")
            
            # Step 4: Convert to PDF
            try:
                pdf_bytes = self.document_processor.doc_to_pdf(merged_doc_bytes)
            except FileNotFoundError as e:
                raise ConversionError(
                    "LibreOffice is not installed. Cannot convert document to PDF."
                )
            except Exception as e:
                raise ConversionError(f"Failed to convert document to PDF: {str(e)}")
            
            # Step 5: Upload to Supabase
            if not file_name:
                # Generate filename with timestamp
                timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                file_name = f"contract_{timestamp}.pdf"
            elif not file_name.endswith('.pdf'):
                file_name = f"{file_name}.pdf"
            
            try:
                upload_result = await self.storage_client.upload_pdf(
                    file_bytes=pdf_bytes,
                    file_name=file_name,
                    user_id=user_id
                )
            except SupabaseUploadError:
                # Re-raise Supabase errors as-is
                raise
            except Exception as e:
                raise SupabaseUploadError(f"Unexpected error during upload: {str(e)}")
            
            # Calculate processing time
            processing_time_ms = (time.time() - start_time) * 1000
            
            # Return result
            return MailMergeResult(
                pdf_bytes=pdf_bytes,
                supabase_upload=upload_result,
                placeholders_replaced=len(placeholders),
                processing_time_ms=processing_time_ms
            )
            
        finally:
            # Step 6: Cleanup temporary files
            self._cleanup_temp_files(temp_files)
    
    def extract_placeholders(self, template_bytes: bytes) -> list[str]:
        """
        Extract list of placeholders from a document template.
        
        Placeholders are in the format {{variable_name}}.
        
        Args:
            template_bytes: Bytes of the .doc/.docx template
            
        Returns:
            List of unique placeholder names (without {{ }})
            
        Raises:
            InvalidTemplateError: If template cannot be read
        """
        import re
        from docx import Document
        
        # Create temporary file for the template
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(template_bytes)
            temp_file_path = temp_file.name
        
        try:
            # Load document
            doc = Document(temp_file_path)
            
            # Collect all text from document
            all_text = []
            
            # Text from paragraphs
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            # Text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            all_text.append(paragraph.text)
            
            # Text from headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    all_text.append(paragraph.text)
                for paragraph in section.footer.paragraphs:
                    all_text.append(paragraph.text)
            
            # Combine all text
            full_text = '\n'.join(all_text)
            
            # Extract placeholders using regex
            # Pattern: {{variable_name}} where variable_name can contain letters, numbers, underscores
            pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'
            matches = re.findall(pattern, full_text)
            
            # Return unique placeholders
            return list(set(matches))
            
        except Exception as e:
            raise InvalidTemplateError(f"Failed to read template: {str(e)}")
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    def validate_merge_data(
        self,
        placeholders: list[str],
        merge_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Validate that merge_data contains all required placeholders.
        
        Args:
            placeholders: List of placeholder names from template
            merge_data: Dictionary with merge data
            
        Returns:
            Dictionary with validation result:
            {
                "is_valid": bool,
                "missing_fields": list[str],
                "extra_fields": list[str]
            }
        """
        merge_data_keys = set(merge_data.keys())
        placeholder_set = set(placeholders)
        
        missing_fields = list(placeholder_set - merge_data_keys)
        extra_fields = list(merge_data_keys - placeholder_set)
        
        return {
            "is_valid": len(missing_fields) == 0,
            "missing_fields": sorted(missing_fields),
            "extra_fields": sorted(extra_fields)
        }
    
    def _cleanup_temp_files(self, temp_files: list[str]):
        """
        Clean up temporary files.
        
        Args:
            temp_files: List of file paths to delete
        """
        for file_path in temp_files:
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
            except Exception:
                # Silently ignore cleanup errors
                pass
